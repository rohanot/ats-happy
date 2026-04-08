import os
import sys
import json
import re
import argparse
import docx
from typing import TypedDict, List, Dict, Any

from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver
from langchain_ollama import ChatOllama
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_groq import ChatGroq
from langchain_core.callbacks import StreamingStdOutCallbackHandler
from dotenv import load_dotenv

# Config
ENGINE = "gemini"
OLLAMA_MODEL = "gemma"

# We use a simple flat-file knowledge base for memory efficiency with small models
KNOWLEDGE_BASE_FILE = "ats_insights.txt"

class ATSState(TypedDict):
    doc_path: str
    resume_text: str
    original_text: str
    jd_text: str
    score: int
    critique: str
    raw_optimizations: str
    changes_applied: List[Dict[str, str]]
    eval_report: str
    iterations: int
    max_iterations: int

def init_llm():
    """Initialize the LLM based on global configuration"""
    if ENGINE == "gemini":
        load_dotenv()
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            print("Error: Please set GEMINI_API_KEY in your .env file.")
            sys.exit(1)
        print("[System] Connecting to Gemini API...")
        return ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.1, streaming=True, callbacks=[StreamingStdOutCallbackHandler()])
    elif ENGINE == "groq":
        load_dotenv()
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            print("Error: Please set GROQ_API_KEY in your .env file (Get one free at console.groq.com)")
            sys.exit(1)
        print(f"[System] Connecting to Groq Cloud (Model: {OLLAMA_MODEL})...")
        return ChatGroq(model=OLLAMA_MODEL, temperature=0.1, streaming=True, callbacks=[StreamingStdOutCallbackHandler()])
    else:
        print(f"[System] Connecting to local Ollama (Model: {OLLAMA_MODEL})...")
        return ChatOllama(model=OLLAMA_MODEL, temperature=0.1, num_ctx=4096, streaming=True, callbacks=[StreamingStdOutCallbackHandler()])

def extract_text_from_docx(doc_path: str) -> tuple[str, Any]:
    doc = docx.Document(doc_path)
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip(): paragraphs.append(p.text.strip())
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip(): paragraphs.append(p.text.strip())
                        
    return "\n".join(paragraphs), doc

def load_insights():
    if os.path.exists(KNOWLEDGE_BASE_FILE):
        with open(KNOWLEDGE_BASE_FILE, "r", encoding="utf-8") as f:
            return f.read()
    return "No past insights yet."

def log_insight(insight):
    with open(KNOWLEDGE_BASE_FILE, "a", encoding="utf-8") as f:
        f.write(f"- {insight}\n")

# --- GRAPH NODES ---

def analyst_node(state: ATSState):
    print(f"\n[Analyst Node] Evaluating Iteration {state['iterations']}...")
    llm = init_llm()
    insights = load_insights()
    
    prompt = f"""You are a strict ATS (Applicant Tracking System) Analyst.
Evaluate the resume below. Be harsh.
Look at past lessons learned from our knowledge base: {insights}

1. First, provide an ATS score strictly out of 100 on the very first line format: SCORE: <number>
2. Then, provide a detailed critique. List the exact paragraphs that lack quantifiable metrics, strong action verbs, or keywords.

Job Description Context:
{state['jd_text']}

Resume Text:
{state['resume_text']}
"""
    response = llm.invoke(prompt)
    output = response.content
    
    # Parse score robustly
    score = 0
    match = re.search(r'SCORE:\s*(\d+)', output, re.IGNORECASE)
    if match:
        score = int(match.group(1))
    else:
        # Fallback regex
        match = re.search(r'\b(\d{1,3})\s*/\s*100', output)
        if match: score = int(match.group(1))
            
    print(f"-> Analyst calculated ATS Score: {score}/100")
    
    return {"score": score, "critique": output, "iterations": state["iterations"] + 1}

def strategist_node(state: ATSState):
    print(f"\n[Strategist Node] Formulating rewrites...")
    llm = init_llm()
    
    prompt = f"""You are an Executive Resume Strategist.
Read the ATS Analyst's critique:
{state['critique']}

Now read the resume text:
{state['resume_text']}

Provide specific textual rewrites for the weak paragraphs identified. 
CRITICAL RULE: DO NOT FABRICATE DATA. DO NOT MAKE UP NUMBERS OR EXPERIENCES. Only rephrase to sound more impactful and action-oriented. You MUST include the exact old paragraph so we can replace it.

You MUST respond strictly in the following JSON format:
[{{
    "old_text": "<exact entire original paragraph>",
    "new_text": "<enhanced paragraph>"
}}]
DO NOT ADD ANY OTHER TEXT EXCEPT THE JSON.
"""
    response = llm.invoke(prompt)
    print("-> Strategist generated optimization plan.")
    return {"raw_optimizations": response.content}

def applier_node(state: ATSState):
    print(f"\n[Applier Node] Parsing AI output and updating layout...")
    raw_text = state['raw_optimizations']
    
    # Robust JSON extraction for small models
    optimizations = []
    match = re.search(r'\[\s*\{.*?\}\s*\]', raw_text, re.DOTALL)
    if match:
        try:
            optimizations = json.loads(match.group(0))
        except Exception as e:
            print(f"Error parsing JSON: {e}")
            
    if not optimizations:
        print("[Applier Node] WARNING: Failed to extract valid JSON from Strategist output. Skipping changes.")
        return {"changes_applied": []}
        
    doc_path = state['doc_path']
    _, doc = extract_text_from_docx(doc_path)
    
    changes = []
    # Note: We append to previous changes if this is iteration > 1
    previous_changes = state.get("changes_applied", [])
    
    for opt in optimizations:
        old_val = opt.get("old_text", "").strip()
        new_val = opt.get("new_text", "").strip()
        if not old_val or not new_val or old_val == new_val:
            continue
            
        replaced = False
        for p in doc.paragraphs:
            if p.text.strip() == old_val:
                p.text = new_val
                replaced = True
                break
                
        if not replaced:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.strip() == old_val:
                                p.text = new_val
                                replaced = True
                                break
                        if replaced: break
                    if replaced: break
                if replaced: break
                
        if replaced:
            changes.append({"old": old_val, "new": new_val})
            
    if changes:
        # Save temporary checkpoint of doc to extract fresh text
        temp_path = doc_path.replace(".docx", "_temp.docx")
        doc.save(temp_path)
        new_text, _ = extract_text_from_docx(temp_path)
        os.remove(temp_path)
        
        # Save real internal state
        doc.save(doc_path)
        
        previous_changes.extend(changes)
        print(f"-> Applied {len(changes)} safe structural inline updates to document.")
        return {"resume_text": new_text, "changes_applied": previous_changes}
    else:
        print("-> No exact text matches found. The agent hallucinated old_text.")
        return {"changes_applied": previous_changes}

def eval_node(state: ATSState):
    print(f"\n[Eval Node] Auditing the AI Pipeline's Work...")
    llm = init_llm()
    
    changes = state.get("changes_applied", [])
    if not changes:
        print("-> No changes applied. Skipping eval.")
        return {"eval_report": "No changes were applied."}
        
    changes_str = "\n".join([f"Old: {c['old']}\nNew: {c['new']}" for c in changes])
    
    prompt = f"""You are the Master Alignment Evaluator.
Your job is to evaluate if the AI pipeline successfully optimized the resume without hallucinating facts.
Look at the changes made:
{changes_str}

Evaluate based on:
1. Did it add quantifiable metrics that did not exist in the old text? (If yes, this is a hallucination! Bad).
2. Did it improve action verbs?
3. Did it maintain the core truth of the resume?

Provide a very short 2-3 sentence report. 
End your output with exactly one "LESSON LEARNED" summarizing what to do differently next time to improve.
Format:
REPORT: ...
LESSON LEARNED: ...
"""
    response = llm.invoke(prompt)
    output = response.content
    print("-> Eval Report Generated.")
    
    # Extract lesson for knowledge base
    match = re.search(r'LESSON LEARNED:\s*(.*)', output, re.IGNORECASE | re.DOTALL)
    if match:
        insight = match.group(1).strip()
        log_insight(insight)
        print(f"-> Logged new insight to knowledge base: {insight}")
        
    return {"eval_report": output}

def route_evaluation(state: ATSState) -> str:
    if state["score"] >= 90:
        print("\n[ROUTER] Target score 90+ achieved! Routing to EVAL.")
        return "eval"
    if state["iterations"] > state["max_iterations"]:
        print(f"\n[ROUTER] Max iterations ({state['max_iterations']}) reached. Routing to EVAL.")
        return "eval"
    
    # Failsafe: if we applied 0 changes in the last loop, the LLM might be stuck
    # Since we can't easily check previous loop changes natively without extra state, we'll just loop.
    print(f"\n[ROUTER] Score {state['score']} < 90. Routing to Strategist for more rewrites.")
    return "strategist"

# --- BUILD GRAPH ---
def build_graph():
    workflow = StateGraph(ATSState)
    
    # Add nodes
    workflow.add_node("analyst", analyst_node)
    workflow.add_node("strategist", strategist_node)
    workflow.add_node("applier", applier_node)
    workflow.add_node("eval", eval_node)
    
    # Add edges
    workflow.add_edge(START, "analyst")
    # Analyst routes conditionally to strategize or eval
    workflow.add_conditional_edges("analyst", route_evaluation, {"eval": "eval", "strategist": "strategist"})
    # Strategist routes to applier
    workflow.add_edge("strategist", "applier")
    # Applier routes back to analyst for re-scoring
    workflow.add_edge("applier", "analyst")
    # Eval ends the graph
    workflow.add_edge("eval", END)
    
    # We use a memory checkpoint
    checkpoint = MemorySaver()
    return workflow.compile(checkpointer=checkpoint)

def main():
    parser = argparse.ArgumentParser(description="LangGraph ATS Optimizer")
    parser.add_argument("resume", help="Path to .docx file")
    parser.add_argument("--jd", help="Path to JD txt", default="")
    parser.add_argument("--out", help="Output .docx path", default="optimized_resume.docx")
    parser.add_argument("--engine", choices=["gemini", "ollama", "groq"], default="gemini", help="Choose AI backend. Default: gemini")
    parser.add_argument("--model", type=str, default="gemma", help="Specific ollama model to run if engine=ollama")
    args = parser.parse_args()
    
    global ENGINE, OLLAMA_MODEL
    ENGINE = args.engine
    OLLAMA_MODEL = args.model
    
    if not os.path.exists(args.resume):
        print(f"File not found: {args.resume}")
        sys.exit(1)
        
    jd_text = "No JD provided."
    if args.jd and os.path.exists(args.jd):
        with open(args.jd, "r") as f:
            jd_text = f.read()
            
    # For a safe test, we copy the resume to the out path immediately and modify it there
    import shutil
    shutil.copy(args.resume, args.out)

    resume_text, _ = extract_text_from_docx(args.out)
    
    app = build_graph()
    
    print("\nStarting LangGraph State Machine...")
    
    initial_state = {
        "doc_path": args.out,
        "resume_text": resume_text,
        "original_text": resume_text,
        "jd_text": jd_text,
        "score": 0,
        "critique": "",
        "raw_optimizations": "",
        "changes_applied": [],
        "eval_report": "",
        "iterations": 1,
        "max_iterations": 3
    }
    
    config = {"configurable": {"thread_id": "ats_thread_1"}}
    
    # Run the graph
    final_state = app.invoke(initial_state, config=config)
    
    print("\n" + "="*40)
    print("PIPELINE COMPLETED")
    print("="*40)
    print(f"Final ATS Score: {final_state['score']}")
    print(f"Total Iterations: {final_state['iterations'] - 1}")
    print(f"\nFinal Eval Report:\n{final_state.get('eval_report', 'N/A')}")
    print("\nSaved fully formatted document to:", args.out)

if __name__ == "__main__":
    main()
